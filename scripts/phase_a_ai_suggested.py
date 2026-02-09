"""
Phase A: YAML-driven ingestion -> cleaning -> chunking -> JSONL outputs.

What this version adds (Option 1):
  - type: folder expansion (glob -> many file sources)
  - real PDF extraction (pypdf)
  - real DOCX extraction (python-docx)
  - ♨️ NEW: Markdown file extraction
  - ♨️ NEW: HTML file extraction

You can now add PDFs/DOCX/MD/HTML as folder entries in data_sources.yaml and run downstream.
"""

import argparse
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml
# Conditional import for pypdf and docx to avoid installation issues if not needed
try:
    from pypdf import PdfReader # type: ignore
except ImportError:
    print("pypdf not installed. PDF loading will be skipped.")
    PdfReader = None

try:
    from docx import Document as DocxDocument # type: ignore
except ImportError:
    print("python-docx not installed. DOCX loading will be skipped.")
    DocxDocument = None

try:
    import requests # type: ignore
except ImportError:
    print("requests not installed. URL loading will be skipped.")
    requests = None

try:
    from bs4 import BeautifulSoup # type: ignore
except ImportError:
    print("BeautifulSoup4 not installed. HTML loading will be degraded/skipped.")
    BeautifulSoup = None

import warnings
# Suppress specific UserWarning from BeautifulSoup if present
warnings.filterwarnings("ignore", category=UserWarning, module='bs4')


# --- Data models ---
@dataclass
class Document:
    id: str
    title: str
    source_type: str   # file/url
    source: str        # abs path or url
    text: str
    meta: Dict[str, Any]

@dataclass
class Chunk:
    doc_id: str
    chunk_id: str
    text: str
    meta: Dict[str, Any]


# --- Cleaning ---
def normalize_text(text: str) -> str:
    """
    Normalizes text by replacing various whitespace characters and
    removing excessive blank lines.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u200b", "").replace("\ufeff", "") # Zero-width space, BOM
    text = text.replace("\t", " ")
    # Replace multiple spaces with a single space
    text = re.sub(r"[ ]+", " ", text)

    lines = [ln.strip() for ln in text.split("\n")]
    out: List[str] = []
    blank = 0
    for ln in lines:
        if ln == "":
            blank += 1
            if blank <= 1: # Allow only one blank line
                out.append("")
        else:
            blank = 0
            out.append(ln)

    return "\n".join(out).strip()

def slugify(s: str) -> str:
    """Converts a string to a slug (lowercase, alphanumeric, hyphen-separated)."""
    s = s.lower().strip()
    s = re.sub(r"[^a-z0-9\s-]", "", s)
    s = re.sub(r"[\s_-]+", "_", s)
    return s.strip("_")


# --- Folder expansion ---
def expand_folder_sources(project_root: Path, sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Converts YAML entries like:
      - type: folder, path: ..., glob: "*.pdf", format: pdf
    into MANY entries of:
      - type: file, path: <each matched file>, format: pdf
    """
    expanded: List[Dict[str, Any]] = []

    for s in sources:
        stype = s.get("type")
        if stype != "folder":
            expanded.append(s)
            continue

        folder_path = s.get("path")
        pattern = s.get("glob")
        fmt = s.get("format")

        if not folder_path or not pattern:
            raise SystemExit(f"Folder source '{s.get('id')}' must include 'path' and 'glob'.")

        folder_abs = (project_root / folder_path).resolve()
        if not folder_abs.exists():
            raise SystemExit(f"Folder path does not exist: {folder_abs}")

        matches = sorted(folder_abs.glob(pattern))
        if not matches:
            print(f"⚠️ No files matched for folder source '{s.get('id')}' at {folder_abs} with glob '{pattern}'")

        base_id = s.get("id", "folder")
        base_title = s.get("title", base_id)
        tags = s.get("tags", [])
        loader_options = s.get("loader_options", {})
        meta_extra = {k: v for k, v in s.items() if k not in {"id", "type", "path", "glob", "format", "title", "tags", "loader_options"}}

        for fp in matches:
            if fp.is_dir():
                continue

            # Build a stable derived ID: baseid__filename_slug
            derived_id = f"{base_id}__{slugify(fp.stem)}"

            expanded.append({
                **meta_extra, # Include any extra metadata from the original source
                "id": derived_id,
                "type": "file",
                "format": fmt, # pdf/docx/txt/...
                "path": str(fp.relative_to(project_root)), # Keep YAML-style relative path
                "title": f"{base_title}: {fp.name}",
                "tags": tags,
                "loader_options": loader_options,
            })
    return expanded


# --- File loaders by format ---
def load_txt_like(path: Path) -> str:
    """Loads plain text files (including .md files as raw text for now)."""
    return path.read_text(encoding="utf-8", errors="replace")

def load_pdf(path: Path) -> str:
    """
    Real PDF text extraction. Uses pypdf.
    """
    if PdfReader is None:
        raise ImportError("pypdf is not installed. Cannot load PDF files. Please run 'pip install pypdf'.")
    
    reader = PdfReader(str(path))
    parts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text:
            parts.append(f"[PAGE {i}]\n{text}")
    return "\n\n".join(parts)

def load_docx(path: Path) -> str:
    """
    Real DOCX text extraction. Uses python-docx.
    """
    if DocxDocument is None:
        raise ImportError("python-docx is not installed. Cannot load DOCX files. Please run 'pip install python-docx'.")
    
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)

def load_html(path: Path) -> str:
    """
    HTML text extraction using BeautifulSoup to get clean text.
    Removes script, style elements, and extracts text content.
    """
    if BeautifulSoup is None:
        raise ImportError("BeautifulSoup4 is not installed. Cannot parse HTML files. Please run 'pip install beautifulsoup4'.")
    
    html_content = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html_content, 'html.parser')

    # Remove script and style elements
    for script_or_style in soup(["script", "style"]):
        script_or_style.decompose()

    # Get text
    text = soup.get_text()

    # Break into lines and remove leading/trailing space on each
    lines = (line.strip() for line in text.splitlines())
    # Break multi-headlines into a line each
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    # Drop blank lines
    text = '\n'.join(chunk for chunk in chunks if chunk)
    return text


def load_file_by_format(project_root: Path, rel_path: str, fmt: str) -> Tuple[str, str]:
    """Dynamically calls the correct loader based on file format."""
    p = (project_root / rel_path).resolve()
    if not p.exists():
        raise SystemExit(f"File not found: {p}")
    
    fmt = (fmt or p.suffix.lstrip('.')).lower() # Use suffix if format not specified

    raw_text: str
    if fmt in {"txt", "md", "markdown"}: # Treat markdown as plain text for now
        raw_text = load_txt_like(p)
    elif fmt == "pdf":
        raw_text = load_pdf(p)
    elif fmt == "docx":
        raw_text = load_docx(p)
    elif fmt == "html":
        raw_text = load_html(p)
    else:
        raise SystemExit(f"Unsupported file format '{fmt}' for file: {p.name}")
    
    return raw_text, str(p)

def load_url(url: str, timeout: int = 30) -> Tuple[str, str]:
    """Loads text content from a URL."""
    if requests is None:
        raise ImportError("requests is not installed. Cannot load URL content. Please run 'pip install requests'.")
    
    try:
        r = requests.get(url, timeout=timeout)
        r.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
        # Attempt to parse HTML to get cleaner content, fall back to raw text
        if 'text/html' in r.headers.get('Content-Type', ''):
            if BeautifulSoup:
                soup = BeautifulSoup(r.text, 'html.parser')
                for script_or_style in soup(["script", "style"]):
                    script_or_style.decompose()
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                content = '\n'.join(chunk for chunk in chunks if chunk)
            else:
                content = r.text # Fallback if BeautifulSoup not available
        else:
            content = r.text
        return content, url
    except requests.exceptions.RequestException as e:
        raise SystemExit(f"Error loading URL {url}: {e}")


# --- Chunkers (keep simple for Phase A; expand later for Phase B/C) ---
def chunk_char_window(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Simple character window chunking with overlap."""
    if chunk_size <= 0:
        return [text]

    out: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + chunk_size)
        out.append(text[i:j])
        if j == n:
            break
        i = max(0, j - overlap)
    return out

def chunk_recursive_langchain(text: str, chunk_size: int, overlap: int) -> Tuple[List[str], str]:
    """
    Uses Langchain's RecursiveCharacterTextSplitter for more sophisticated chunking.
    Requires langchain to be installed.
    Adapted from 17th Jan and 31st Jan notebooks.
    """
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        return [text], "Langchain not installed - falling back to char_window"

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        is_separator_regex=False, # Use default separators
    )
    chunks = splitter.split_text(text)
    return chunks, "RecursiveCharacterTextSplitter"


# --- Main processing logic ---
def main():
    parser = argparse.ArgumentParser(description="Phase A: Ingest, Clean, Chunk documents")
    parser.add_argument("--data-sources", type=str, default="data_sources.yaml",
                        help="Path to the YAML file defining data sources.")
    parser.add_argument("--output-docs", type=str, default="outputs/documents_all.jsonl",
                        help="Output path for processed documents (JSONL).")
    parser.add_argument("--output-chunks", type=str, default="outputs/chunks_all_recursive.jsonl",
                        help="Output path for generated chunks (JSONL).")
    parser.add_argument("--project-root", type=str, default=".",
                        help="Project root directory for resolving relative paths.")
    
    args = parser.parse_args()

    project_root = Path(args.project_root).resolve()
    output_docs_path = Path(args.output_docs)
    output_chunks_path = Path(args.output_chunks)

    output_docs_path.parent.mkdir(parents=True, exist_ok=True)
    output_chunks_path.parent.mkdir(parents=True, exist_ok=True)

    # Load data sources from YAML
    with open(args.data_sources, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)
    
    initial_sources = config.get("sources", [])
    
    # Expand folder sources (e.g., glob patterns for PDFs)
    sources = expand_folder_sources(project_root, initial_sources)
    
    all_documents: List[Document] = []
    all_chunks: List[Chunk] = []

    rag_config = config.get('rag', {})
    chunk_size = rag_config.get('chunk_size', 500)
    chunk_overlap = rag_config.get('chunk_overlap', 50)

    print(f"--- Processing {len(sources)} data sources ---")

    for source_entry in sources:
        doc_id = source_entry["id"]
        source_type = source_entry["type"]
        doc_title = source_entry.get("title", doc_id)
        
        raw_content: str = ""
        actual_source_path: str = ""
        
        if source_type == "file":
            format_type = source_entry.get("format", "")
            file_path = source_entry["path"]
            try:
                raw_content, actual_source_path = load_file_by_format(project_root, file_path, format_type)
            except SystemExit as e:
                print(f"Error processing file '{file_path}': {e}")
                continue # Skip to next source
        elif source_type == "url":
            url = source_entry["url"]
            try:
                raw_content, actual_source_path = load_url(url)
            except SystemExit as e:
                print(f"Error processing URL '{url}': {e}")
                continue # Skip to next source
        else:
            print(f"Unsupported source type '{source_type}' for ID '{doc_id}'. Skipping.")
            continue

        normalized_text = normalize_text(raw_content)

        # Create Document object
        document_meta = {k: v for k, v in source_entry.items() if k not in {"id", "type", "path", "glob", "format", "title", "url"}}
        current_document = Document(
            id=doc_id,
            title=doc_title,
            source_type=source_type,
            source=actual_source_path,
            text=normalized_text,
            meta=document_meta
        )
        all_documents.append(current_document)

        # Chunk the document
        chunks, chunking_method = chunk_recursive_langchain(normalized_text, chunk_size, chunk_overlap)
        # Fallback if langchain is not present
        if chunking_method == "Langchain not installed - falling back to char_window":
            print(f"Warning: Langchain not installed. Falling back to char_window chunking for '{doc_id}'.")
            chunks = chunk_char_window(normalized_text, chunk_size, chunk_overlap)

        for i, chk_text in enumerate(chunks):
            chunk_obj = Chunk(
                doc_id=doc_id,
                chunk_id=f"{doc_id}-{i:04d}",
                text=chk_text,
                meta={
                    "doc_id": doc_id,
                    "title": doc_title,
                    "source_type": source_type,
                    "source": actual_source_path,
                    "chunk_idx": i,
                    "chunk_size": len(chk_text),
                    "chunking_method": chunking_method,
                    **document_meta # Include any extra metadata from the original source
                }
            )
            all_chunks.append(chunk_obj)
        
        print(f"Processed '{doc_id}' ({source_type}): {len(chunks)} chunks.")

    # Save documents and chunks to JSONL files
    with open(output_docs_path, "w", encoding="utf-8") as f:
        for doc in all_documents:
            f.write(json.dumps(asdict(doc)) + "\n")
    print(f"Saved {len(all_documents)} documents to {output_docs_path}")

    with open(output_chunks_path, "w", encoding="utf-8") as f:
        for chunk in all_chunks:
            f.write(json.dumps(asdict(chunk)) + "\n")
    print(f"Saved {len(all_chunks)} chunks to {output_chunks_path}")

if __name__ == "__main__":
    main()
