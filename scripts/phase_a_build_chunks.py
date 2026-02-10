"""
Phase A: YAML-driven ingestion -> cleaning -> chunking -> JSONL outputs.

What this version adds (Option 1):
✅ type: folder expansion (glob -> many file sources)
✅ real PDF extraction (pypdf)
✅ real DOCX extraction (python-docx)
✅ CSV/TSV extraction (each row becomes a text chunk with headers).
✅ Excel (.xlsx, .xls) extraction (each row/sheet becomes a text chunk with headers).
✅ Markdown and HTML extraction
✅ Chunking strategies:
    # RecursiveCharacterTextSplitter for intelligent splitting based on configurable separators.
    # Sentence-based pre-processing for better contextual chunks.
    # Metadata enrichment: Page numbers, section hints, filename, etc., stored with chunks.
    # Error handling and warnings for missing dependencies.


You can now add PDFs/DOCX as folder entries in data_sources.yaml and run downstream.
"""

import argparse
import json
import re
from dataclasses import dataclass, asdict, field
from pathlib import Path
from time import time
from typing import Any, Dict, List, Optional, Tuple, Callable, Type # Added Type for type hinting chunker class
import hashlib

import yaml
import warnings
import logging
import os
import sys
import time

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# --- Internal RAG system imports ---
# Adjust import paths as necessary based on your project structure

# Import Settings and load_settings
from config import Settings, load_settings

# Import the chunker registry and specific chunker classes
# Ensure these chunkers are imported so they register themselves in CHUNK_STRATEGY_REGISTRY
try:
    from scripts.rag.components.chunkers import (
        Chunk, # Your Chunk dataclass defined in chunkers.py
        CHUNK_STRATEGY_REGISTRY,
        BaseChunkingStrategy,
        RecursiveCharacterChunking,
        SentenceChunking,
        MarkdownHeaderChunking,
        HTMLSectionChunking,
        PropositionChunking, # <<< Added PropositionChunking import
        ParentDocumentChunking # <<< Added ParentDocumentChunking import
    )
except ImportError as e:
    logger.error(f"Error importing chunking components from scripts.rag.components.chunkers: {e}. "
                 "Ensure scripts/rag/components/chunkers.py is correctly set up and dependencies are met.")
    # Set fallbacks so the script can still run in a degraded mode if necessary
    Chunk = None
    CHUNK_STRATEGY_REGISTRY = {}
    BaseChunkingStrategy = type('BaseChunkingStrategy', (object,), {'__init__': lambda s, *a, **k: None, 'split_text': lambda s, *a, **k: []}) # Dummy Base class
    RecursiveCharacterChunking = None
    SentenceChunking = None
    MarkdownHeaderChunking = None
    HTMLSectionChunking = None
    PropositionChunking = None
    ParentDocumentChunking = None 

# --- Rest of the imports remain the same ---
# PDF
try:
    from pypdf import PdfReader  # type: ignore
except ImportError:
    PdfReader = None
    logger.warning("pypdf not installed. PDF loading disabled. Install: pip install pypdf")

# DOCX
try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed. DOCX loading disabled. Install: pip install python-docx")

# Requests (for URLs)
try:
    import requests  # type: ignore
except ImportError:
    requests = None
    logger.warning("requests not installed. URL loading disabled. Install: pip install requests")

# HTML parsing (BeautifulSoup4)
try:
    from bs4 import BeautifulSoup  # type: ignore
    warnings.filterwarnings("ignore", category=UserWarning, module="bs4")
except ImportError:
    BeautifulSoup = None
    logger.warning("beautifulsoup4 not installed. HTML parsing will be raw. Install: pip install beautifulsoup4")

# Markdown parsing (better)
try:
    from markdown_it import MarkdownIt  # type: ignore
    from mdit_py_plugins.front_matter import front_matter_plugin  # type: ignore
except ImportError:
    MarkdownIt = None
    front_matter_plugin = None
    logger.warning("markdown-it-py / mdit-py-plugins not installed. Markdown parsing will be raw. "
                   "Install: pip install markdown-it-py mdit-py-plugins")

# CSV / Excel (pandas)
try:
    import pandas as pd  # type: ignore
except ImportError:
    pd = None
    logger.warning("pandas not installed. CSV/Excel loading disabled. Install: pip install pandas")

# Openpyxl (for .xlsx support in pandas)
try:
    import openpyxl  # type: ignore
except ImportError:
    openpyxl = None
    logger.warning("openpyxl not installed. Excel (.xlsx) loading may fail. Install: pip install openpyxl")

# NLTK sentence tokenization (As discussed, kept for now, but will be less relevant for LangChain SentenceChunking)
try:
    import nltk  # type: ignore
    try:
        nltk.data.find("tokenizers/punkt")
    except LookupError:
        logger.info("Downloading NLTK punkt tokenizer...")
        nltk.download("punkt", quiet=True)
    sentence_tokenizer = nltk.sent_tokenize # This will still be available but likely unused by the new system
except Exception:
    nltk = None
    sentence_tokenizer = lambda text: text.split(". ") # Basic fallback
    logger.warning("NLTK not available or punkt missing. Sentence operations will be basic.")

# -----------------------------
# Helpers 
# -----------------------------
def get_config_hash(config_dict: Dict[str, Any]) -> str:
    """Generates a stable hash for a configuration dictionary."""
    # Convert dict to a sorted JSON string for consistent hashing
    sorted_json = json.dumps(config_dict, sort_keys=True, ensure_ascii=False)
    return hashlib.md5(sorted_json.encode('utf-8')).hexdigest()

# --- Data Models ---
# This Document dataclass remains local to phase_a_build_chunks for ingestion processing
# The Chunk dataclass is now imported from src.rag.components.chunkers.py
@dataclass
class Document:
    """
    Represents a full document before chunking.
    id: Unique identifier for the document
    title: Human-readable title
    source_type: file/url
    source: absolute path or URL
    text: Full document text
    meta: Other metadata (tags, loader_options, etc.)
    """
    id: str
    title: str
    source_type: str  # file/url
    source: str       # abs path or url (absolute path for files, URL for urls)
    text: str # Full document content after extraction and cleaning
    meta: Dict[str, Any] = field(default_factory=dict) # Other metadata (tags, loader_options, etc.)


# scripts/phase_a_build_chunks.py (continued from Step 1)

def load_config(config_path: Path) -> Dict[str, Any]:
    """Loads configuration from a YAML file."""
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def load_data_sources(data_sources_path: Path) -> List[Dict[str, Any]]:
    """Loads data source definitions from a YAML file."""
    with open(data_sources_path, 'r', encoding='utf-8') as f:
        # Assumes 'sources' is the top-level key as per your data_sources.yaml
        return yaml.safe_load(f)["sources"]

# -----------------------------
# Cleaning
# -----------------------------
def normalize_text(text: str) -> str:
    """
    Normalizes text by replacing various whitespace characters,
    removing excessive blank lines, and stripping leading/trailing whitespace.
    """
    if not isinstance(text, str):
        return "" # Ensure we always return a string
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Common zero-width and non-breaking spaces
    text = text.replace("\u200b", "").replace("\ufeff", "").replace("\xa0", " ")
    text = text.replace("\t", " ") # Replace tabs with single space
    text = re.sub(r"[ ]+", " ", text) # Replace multiple spaces with a single space

    lines = [ln.strip() for ln in text.split("\n")]
    out: List[str] = []
    blank_count = 0
    for ln in lines:
        if not ln:  # If line is empty after stripping
            blank_count += 1
            if blank_count <= 1: # Allow only one blank line
                out.append("")
        else:
            blank_count = 0
            out.append(ln)

    return "\n".join(out).strip()

def slugify(s: str) -> str:
    """Converts a string to a slug (lowercase, alphanumeric, hyphen-separated)."""
    s = s.lower().strip()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_-]+", "_", s) # Use underscore per your existing style
    return s.strip("_")

def safe_meta(meta: Dict[str, Any]) -> Dict[str, Any]:
    """
    Converts metadata values to types compatible with ChromaDB (or other vector stores).
    Chroma metadata values must be: str, int, float, bool, or None.
    Lists and dicts should be JSON-serialized.
    """
    out_meta = {}
    for k, v in meta.items():
        if v is None or isinstance(v, (str, int, float, bool)):
            out_meta[k] = v
        elif isinstance(v, (list, tuple, set, dict)):
            # JSON-serialize complex structures
            try:
                out_meta[k] = json.dumps(v, ensure_ascii=False)
            except TypeError:
                logger.warning(f"Could not JSON-serialize metadata key '{k}'. Converting to string.")
                out_meta[k] = str(v)
        else:
            # Fallback for other types
            out_meta[k] = str(v)
    return out_meta

# -----------------------------
# Folder expansion
# -----------------------------
def expand_folder_sources(project_root: Path, sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Converts YAML entries like:
      - type: folder, path: data/corpus, glob: "**/*.pdf", format: pdf
    into MANY entries like:
      - type: file, path: data/corpus/file1.pdf, format: pdf

    Infers format from extension if 'format: auto' or missing.
    """
    expanded: List[Dict[str, Any]] = []

    for s in sources:
        stype = s.get("type")
        if stype != "folder":
            expanded.append(s)
            continue

        folder_path = s.get("path")
        pattern = s.get("glob")
        fmt = s.get("format") # 'pdf', 'docx', 'txt', 'auto', None etc.

        if not folder_path or not pattern:
            logger.error(f"Folder source '{s.get('id','N/A')}' must include 'path' and 'glob'. Skipping.")
            continue

        folder_abs = (project_root / folder_path).resolve()
        if not folder_abs.exists() or not folder_abs.is_dir():
            logger.warning(f"Folder path invalid for '{s.get('id')}': {folder_abs}. Skipping.")
            continue

        matches = sorted(folder_abs.glob(pattern))
        if not matches:
            logger.info(f"No files matched for folder source '{s.get('id')}' with glob '{pattern}'")
            continue

        base_id = s.get("id", slugify(folder_abs.name))
        base_title = s.get("title", base_id.replace("_", " ").title())
        tags = s.get("tags", [])
        loader_options = s.get("loader_options", {})
        chunking_config = s.get("chunking") # <<< Get chunking config from parent source

        # Any other metadata from the source definition is passed through
        meta_extra = {k: v for k, v in s.items() if k not in {"id", "type", "path", "glob", "format", "title", "tags", "loader_options", "chunking"}}
 
        for fp in matches:
            if fp.is_dir(): # Skip subdirectories
                continue

            # Determine format if 'auto'
            file_fmt = fmt
            if fmt == ["auto", None]:
                suffix = fp.suffix.lstrip(".").lower()
                if suffix in ["txt", "pdf", "docx", "md", "html", "csv", "tsv", "xls", "xlsx"]:
                    file_fmt = suffix
                else:
                    logger.warning(f"Could not infer format for {fp.name}. Defaulting to 'txt'.")
                    file_fmt = "txt"

            # Build a stable derived ID: baseid__filename_slug
            # Using relative path hash for IDs for stability
            rel_path_str = str(fp.relative_to(project_root))
            path_hash = hashlib.md5(rel_path_str.encode()).hexdigest()[:8]
            derived_id = f"{base_id}__{path_hash}"

            # Add source entry with absolute path (for loader) but relative_to for storage/metadata
            expanded_entry = {
                **meta_extra, # Include any extra metadata from original source
                "id": derived_id,
                "type": "file",
                "format": file_fmt, # Determined format
                "path": rel_path_str, # Relative path to project root
                "abs_path": str(fp.resolve()), # Absolute path on disk for opening
                "title": f"{base_title}: {fp.name}",
                "tags": tags,
                "loader_options": loader_options,
            }
            if chunking_config: # Pass down source-specific chunking if defined
                expanded_entry["chunking"] = chunking_config
            
            expanded.append(expanded_entry)
            
    return expanded


# -----------------------------
# File loaders by format
# -----------------------------
def load_txt_like(path: Path) -> str:
    """Loads plain text files (including .md if no specific MD parser)."""
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Error reading text file {path}: {e}")
        return ""


def load_pdf(path: Path) -> str:
    """Reads PDF text using pypdf."""
    # Fixed docstring closing quote
    if PdfReader is None:
        raise ImportError("pypdf missing")
    reader = PdfReader(str(path))
    parts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        txt = page.extract_text()
        if txt:
            parts.append(f"[PAGE {i}]\n{txt}")
    return "\n\n".join(parts)

def load_docx(path: Path) -> str:
    """Reads DOCX text using python-docx."""
    if DocxDocument is None:
        raise ImportError("python-docx not installed. Cannot load DOCX files.")
    
    doc = DocxDocument(path)
    # Extract text from paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Attempt to extract text from tables (simple approach)
    tables_text = []
    for table in doc.tables:
        table_rows = []
        # Each row in the table
        for r_idx, row in enumerate(table.rows):
            # Extract text from cells in the row
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                # Add a prefix to distinguish rows from tables
                tables_text.append(f"Table Row {r_idx+1}: " + " | ".join(row_cells))
            
    all_content = "\n\n".join(paragraphs + tables_text)
    return all_content

def load_html_raw(path: Path) -> str:
    """Loads HTML from a file as raw text."""
    return path.read_text(encoding="utf-8", errors="replace")

def load_html(path: Path) -> str:
    """Reads HTML text using BeautifulSoup to extract visible text."""
    if BeautifulSoup is None:
        logger.warning("BeautifulSoup4 not available. Loading HTML raw.")
        return load_txt_like(path) # Fallback to raw load via load_txt_like
    
    html_content = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html_content, "html.parser")

    # Remove script, style, and other non-text elements (more comprehensive list)
    for element in soup(["script", "style", "header", "footer", "nav", "form", "meta", "link", "img", "svg"]):
        element.extract()
    
    # Get text
    text = soup.get_text(separator="\n", strip=True) # Use newline as separator for better readability

    return text

# Refined load_md_with_frontmatter
def load_md_with_frontmatter(path: Path) -> Tuple[str, Dict[str, Any]]:
    """Loads Markdown, parses front matter, returning clean text and extracted metadata."""
    if MarkdownIt is None or front_matter_plugin is None:
        logger.warning(f"Markdown-it/front_matter_plugin missing for {path}. Loading as plain text.")
        return load_txt_like(path), {} # Return empty dict for metadata and full text as is
    
    text_content = path.read_text(encoding="utf-8", errors="replace")
    
    # --- Manual Front Matter Parsing for robustness and direct YAML control ---
    front_matter_meta: Dict[str, Any] = {}
    cleaned_text = text_content
    
    front_matter_pattern = re.compile(r"^(---|\+\+\+)\s*\n(.*?)\n(---|\+\+\+)\s*\n", re.DOTALL)
    match = front_matter_pattern.match(text_content)
    
    if match:
        yaml_block = match.group(2)
        try:
            front_matter_meta = yaml.safe_load(yaml_block) or {}
            cleaned_text = text_content[match.end():]
        except yaml.YAMLError as yml_err:
            logger.warning(f"Error parsing front matter in {path}: {yml_err}. Content will be treated as plain text.")
            cleaned_text = text_content # Fallback to original content
    
    return cleaned_text.strip(), front_matter_meta # Return stripped text and extracted meta


def load_csv_like(path: Path, sep: str = ',', header: bool = True, loader_options: Optional[Dict[str, Any]] = None) -> Tuple[str, List[str]]:
    """Loads CSV/TSV files, treating each row as a 'pre-chunk' for more structured processing."""
    if pd is None:
        raise ImportError("pandas not installed. Cannot load CSV/TSV.")

    loader_options = loader_options or {}
    try:
        df = pd.read_csv(str(path), sep=sep, **loader_options)
    except Exception as e:
        logger.error(f"Error reading {path} as CSV/TSV: {e}")
        return "", []

    extracted_rows: List[str] = []
    # Include headers for each row for better context
    column_names = df.columns.tolist()
    for _, row in df.iterrows():
        # Using a list comprehension to build row parts, filtering out NaNs
        row_text_parts = []
        for col_name in column_names:
            val = row[col_name]
            if pd.notna(val):
                row_text_parts.append(f"{col_name}: {val}")
        
        if row_text_parts:
            extracted_rows.append("; ".join(row_text_parts))

    # A simple combined text might be useful for some chunkers, but for structured, pre-chunks are key.
    full_text = "\n\n".join(extracted_rows)
    return full_text, extracted_rows


def load_excel(path: Path, loader_options: Optional[Dict[str, Any]] = None) -> Tuple[str, List[str]]:
    """Loads Excel files (xlsx, xls), treating each row from each sheet as a 'pre-chunk'."""
    if pd is None:
        raise ImportError("pandas not installed. Cannot load Excel.")
    # Check openpyxl only for .xlsx, as pandas can use xlrd for .xls
    if path.suffix == ".xlsx" and openpyxl is None:
        raise ImportError("openpyxl not installed. Cannot load .xlsx files.")

    loader_options = loader_options or {}
    all_rows: List[str] = []
    full_text_parts: List[str] = []

    try:
        # read_excel can read all sheets by default
        excel_data = pd.read_excel(str(path), sheet_name=None, **loader_options)
        
        for sheet_name, df in excel_data.items():
            full_text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
            column_names = df.columns.tolist()
            for r_idx, row in df.iterrows():
                row_text_parts = []
                for col_name in column_names:
                    val = row[col_name]
                    if pd.notna(val):
                        row_text_parts.append(f"{col_name}: {val}")
                
                if row_text_parts:
                    row_content = f"Sheet: {sheet_name}, Row {r_idx+1}: " + "; ".join(row_text_parts)
                    all_rows.append(row_content)
                    full_text_parts.append(row_content)
                
    except Exception as e:
        logger.error(f"Error reading Excel file {path}: {e}")
        return "", []

    return "\n\n".join(full_text_parts).strip(), all_rows


def load_url(url: str, timeout: int = 30) -> str:
    """Loads content from a URL."""
    if requests is None:
        raise ImportError("requests not installed. Cannot load URLs.")
    try:
        response = requests.get(url, timeout=timeout)
        response.raise_for_status() # Raise an exception for HTTP errors

        content_type = response.headers.get("Content-Type", "").lower()
        if "text/html" in content_type and BeautifulSoup:
            soup = BeautifulSoup(response.text, 'html.parser')
            # Extract main readable text from HTML, removing common non-content elements
            for element in soup(["script", "style", "head", "footer", "nav", "form", "meta", "link", "img", "svg"]):
                element.extract()
            # Use ' ' as separator for in-line text, newlines for blocks
            text = soup.get_text(separator=' ', strip=True)
            # Replace multiple spaces/newlines with single relevant separator
            text = re.sub(r'\s{2,}', ' ', text).strip() # Compress multiple whitespace
            return text
        else:
            # Fallback for non-HTML or if BS4 is missing
            return response.text
    except requests.exceptions.RequestException as e:
        logger.error(f"Error loading URL {url}: {e}")
        return ""


# This is the dispatcher function that main() will call
def load_file_by_format(
    project_root: Path,
    file_path_relative: str,
    file_format: str,
    loader_options: Optional[Dict[str, Any]] = None
) -> Tuple[str, List[str], str, Dict[str, Any]]: # Added Dict[str, Any] for extra_meta
    """
    Dispatches file loading based on format.
    Returns (full_text, pre_chunks, absolute_source_path, extra_metadata).
    extra_metadata is for things like markdown front matter.
    """
    abs_path = (project_root / file_path_relative).resolve()
    full_text = ""
    pre_chunks: List[str] = []
    extra_meta: Dict[str, Any] = {}

    try:
        if file_format == "txt":
            full_text = load_txt_like(abs_path)
        elif file_format == "pdf":
            full_text = load_pdf(abs_path)
        elif file_format == "docx":
            full_text = load_docx(abs_path)
        # For md and html, return raw content for structural splitters, but clean text for general.
        # MarkdownHeaderTextSplitter and HTMLSectionSplitter take raw content.
        elif file_format == "md":
            # load_md_with_frontmatter returns (cleaned_text, metadata)
            # Use this cleaned_text for generic text chunking, and metadata for meta enrichment
            full_text, extra_meta = load_md_with_frontmatter(abs_path)
            # For structural Markdown chunking, the "splitter" itself would process the raw file directly,
            # but for consistency with this loader, we're providing the text content.
            # The actual MarkdownHeaderTextSplitter works better on raw markdown,
            # so we might need a separate path for that or let the splitter parse.
            # For now, it will work on the 'full_text' as returned.
        elif file_format == "html":
            # load_html returns extracted text, not raw HTML.
            # Same note as MD: For structural HTML chunking, the "splitter" itself would process raw HTML.
            # Here, we pass the extracted text.
            full_text = load_html(abs_path)
        elif file_format == "csv":
            full_text, pre_chunks = load_csv_like(abs_path, sep=',', loader_options=loader_options)
        elif file_format == "tsv":
            full_text, pre_chunks = load_csv_like(abs_path, sep='\t', loader_options=loader_options)
        elif file_format in ["xls", "xlsx"]:
            full_text, pre_chunks = load_excel(abs_path, loader_options=loader_options)
        else:
            logger.warning(f"Unsupported format '{file_format}' for {file_path_relative}. Loading as plain text.")
            full_text = load_txt_like(abs_path)
    except ImportError as ie:
        logger.error(f"Missing dependency for {file_format} file '{abs_path}': {ie}")
        return "", [], str(abs_path), {}
    except Exception as e:
        logger.error(f"Failed to load '{abs_path}' (format: {file_format}): {e}")
        return "", [], str(abs_path), {}

    return full_text, pre_chunks, str(abs_path), extra_meta


# -----------------------------
# Selection helper
# -----------------------------
def select_sources(sources: List[Dict[str, Any]], ids: Optional[List[str]], exclude: Optional[List[str]]) -> List[Dict[str, Any]]:
    wanted = set(ids) if ids else None
    excluded = set(exclude) if exclude else set()

    selected: List[Dict[str, Any]] = []
    for s in sources:
        sid = s.get("id")
        if not sid:
            raise SystemExit("A source is missing required field: id")
        if wanted is not None and sid not in wanted:
            continue
        if sid in excluded:
            continue
        selected.append(s)
    return selected

# --- Main Manifest Management (NEW) ---
MANIFEST_FILE_NAME = "phase_a_manifest.json"

def load_manifest(manifest_path: Path) -> Dict[str, Any]:
    """Loads the processing manifest."""
    if manifest_path.exists():
        with manifest_path.open("r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError as e:
                logger.error(f"Error reading manifest file {manifest_path}: {e}. Starting with empty manifest.")
                return {}
    return {}

def save_manifest(manifest_path: Path, manifest_data: Dict[str, Any]):
    """Saves the processing manifest."""
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    with manifest_path.open("w", encoding="utf-8") as f:
        json.dump(manifest_data, f, indent=2, ensure_ascii=False)
    logger.info(f"Updated manifest saved to {manifest_path}")


# -----------------------------
# Main
# -----------------------------
# In main() before calling load_config
# scripts/phase_a_build_chunks.py (continued from Step 2)

def main():
    parser = argparse.ArgumentParser(
        description="Phase A: Ingest, clean, chunk, and output documents/chunks in JSONL format."
    )
    parser.add_argument(
        "--project-root",
        type=Path,
        default=".",
        help="Root directory of the project (for resolving paths).",
    )
    parser.add_argument(
        "--config-path",
        type=Path,
        default="config.yaml",
        help="Path to the main configuration YAML file.",
    )
    parser.add_argument(
        "--data-sources-path",
        type=Path,
        default="data_sources.yaml",
        help="Path to the data sources YAML file.",
    )
    # The default for output-docs/chunks/parents will be dynamically overridden later.
    parser.add_argument(
        "--output-docs",
        type=Path,
        default="outputs/phase_a_processed_docs.jsonl", # Base default, overridden dynamically
        help="Output path for processed documents (pre-chunking)."
    )
    parser.add_argument(
        "--output-chunks",
        type=Path,
        default="outputs/phase_a_processed_chunks.jsonl", # Base default, overridden dynamically
        help="Output path for processed chunks."
    )
    parser.add_argument(
        "--output-parents",
        type=Path,
        default="outputs/phase_a_processed_parents.jsonl", # Base default, overridden dynamically
        help="Output path for parent documents (larger context units)."
    )
    parser.add_argument(
        "--timeout",
        type=int,
        default=60,
        help="Timeout in seconds for URL requests."
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        help="Set the logging level."
    )
    parser.add_argument(
        "--include-sources",
        nargs='*',
        help="List of source IDs from data_sources.yaml to include (space-separated). "
             "Only these will be processed."
    )
    parser.add_argument(
        "--exclude-sources",
        nargs='*',
        help="List of source IDs from data_sources.yaml to exclude (space-separated)."
    )
    parser.add_argument(
        "--force-reprocess",
        action="store_true",
        help="Force re-processing of all selected sources, even if existing outputs match config."
    )
    # Arguments for dynamic output path (these are for naming the output JSONL files distinctly)
    parser.add_argument(
        "--output-suffix-chunking-strategy", # Renamed for clarity to avoid confusion with Chroma path
        type=str,
        default="default",
        help="Chunking strategy name to embed in output JSONL filenames (e.g., 'recursive_character')."
    )
    parser.add_argument(
        "--output-suffix-doc-type", # Renamed for clarity
        type=str,
        default="docs",
        help="Document type name to embed in output JSONL filenames (e.g., 'reports', 'manuals')."
    )

    args = parser.parse_args()

    logger.setLevel(getattr(logging, args.log_level.upper()))
    
    # --- Load global settings and RAG config ---
    project_root = args.project_root.resolve()
    settings = load_settings() # Load environment variables
    
    full_config = load_config(args.config_path)
    rag_config = full_config.get("rag", {})
    global_chunking_config = rag_config.get("chunking", {})

    logger.info(f"Loaded RAG configuration: {rag_config}")
    logger.info(f"Global chunking configuration: {global_chunking_config}")

    # Determine chunking strategy name for output filenames (from CLI then config default)
    overall_chunking_strategy_name = args.output_suffix_chunking_strategy
    if overall_chunking_strategy_name == "default" and "chunking" in rag_config:
        overall_chunking_strategy_name = rag_config["chunking"].get("strategy", "recursive_character")
    
    # Generate a hash for the overall RAG pipeline configuration ---
    # This captures the state of the entire RAG config (including reranker, MQR, distiller settings).
    # Use full_config.get("rag", {}) because that contains general RAG parameters PLUS
    # the nested 'reranker', 'query_transformer', 'context_distiller' settings.
    overall_rag_pipeline_config_hash = get_config_hash(full_config.get("rag", {}))
    logger.info(f"Overall RAG pipeline config hash: {overall_rag_pipeline_config_hash}")


    # --- Dynamic Output Paths Construction (for JSONL files) ---
    dynamic_suffix_parts = []
    dynamic_suffix_parts.append(model_slug(overall_chunking_strategy_name))
    if args.output_suffix_doc_type: # Using output-suffix-doc-type arg
        dynamic_suffix_parts.append(model_slug(args.output_suffix_doc_type))
    
    dynamic_suffix = "_".join(dynamic_suffix_parts) if dynamic_suffix_parts else "default"

    base_output_dir = project_root / "outputs" # Centralized output folder, relative to project_root

    final_output_docs_path = base_output_dir / Path(f"phase_a_processed_docs_{dynamic_suffix}.jsonl")
    final_output_chunks_path = base_output_dir / Path(f"phase_a_processed_chunks_{dynamic_suffix}.jsonl")
    final_output_parents_path = base_output_dir / Path(f"phase_a_processed_parents_{dynamic_suffix}.jsonl")
    
    base_output_dir.mkdir(parents=True, exist_ok=True) # Ensure outputs directory exists
    
    logger.info(f"Output for documents: {final_output_docs_path}")
    logger.info(f"Output for child chunks: {final_output_chunks_path}")
    logger.info(f"Output for parent chunks: {final_output_parents_path}")
    logger.info(f"Sources to include: {args.include_sources or 'All'}")
    logger.info(f"Sources to exclude: {args.exclude_sources or 'None'}")
    logger.info(f"Force re-process: {args.force_reprocess}")

    # --- Load data sources and apply initial selection filtering ---
    doc_list_raw: List[Dict[str, Any]] = load_data_sources(args.data_sources_path)
    expanded_sources_full = expand_folder_sources(project_root, doc_list_raw)
    sources_to_process_filtered = select_sources(expanded_sources_full, args.include_sources, args.exclude_sources)
    logger.info(f"Selected {len(sources_to_process_filtered)} specific document sources for processing.")

    # Load Manifest (for skipping/reprocessing logic) ---
    manifest_path = base_output_dir / MANIFEST_FILE_NAME
    current_manifest = load_manifest(manifest_path)
    updated_manifest = current_manifest.copy()

    # --- Prepare for comprehensive output information ---
    summary_output = {
        "phase": "Phase A: Ingestion & Chunking",
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "total_sources_in_config": len(doc_list_raw),
        "total_expanded_sources": len(expanded_sources_full),
        "total_sources_selected_for_processing": len(sources_to_process_filtered),
        "output_files": {
            "docs": str(final_output_docs_path.relative_to(project_root)),
            "chunks": str(final_output_chunks_path.relative_to(project_root)),
            "parents": str(final_output_parents_path.relative_to(project_root)),
        },
        "overall_chunking_strategy_used_for_naming": overall_chunking_strategy_name,
        "document_type_for_naming": args.output_suffix_doc_type,
        "rag_pipeline_config_hash_for_run": overall_rag_pipeline_config_hash,
        "processed_sources_details": [],
        "total_documents_loaded": 0,
        "total_chunks_written": 0,
        "total_parent_chunks_written": 0,
    }

    
    # --- Begin processing individual sources ---
    documents: List[Document] = []
    
    for s_conf in sources_to_process_filtered:
        sid = s_conf["id"]
        source_title = s_conf.get("title", sid)
        source_type = s_conf.get("type")
        source_format = s_conf.get("format")
        source_location = s_conf.get("path", s_conf.get("url", "N/A"))

        # Get chunking config specific to this source, if any, else global default
        doc_specific_chunking_config = s_conf.get("chunking", global_chunking_config)
        
        # Hash the configuration for this specific source to check for reprocessing needs
        source_processing_config_hash_dict = {
            "source_id": sid,
            "format": source_format,
            "type": source_type,
            "loader_options": s_conf.get("loader_options", {}),
            "chunking_config": doc_specific_chunking_config
        }
        current_source_config_hash = get_config_hash(source_processing_config_hash_dict) # Generate hash for current config
        
        doc_summary_entry = {
            "id": sid,
            "title": source_title,
            "location": source_location,
            "format": source_format,
            "chunking_strategy_planned": doc_specific_chunking_config.get("strategy"),
            "source_processing_config_hash": current_source_config_hash,
            "pipeline_config_hash": overall_rag_pipeline_config_hash,
            "status": "skipped_reprocess",
            "message": "Skipped (already processed with same config)",
            "chunks_generated": 0,
            "parents_generated": 0
        }
        
        summary_output["processed_sources_details"].append(doc_summary_entry)

        # --- REPROCESSING / CACHING CHECK ---
        manifest_entry = current_manifest.get(sid)
        if (not args.force_reprocess and
            manifest_entry and
            manifest_entry.get("source_processing_config_hash") == current_source_config_hash and
            manifest_entry.get("pipeline_config_hash") == overall_rag_pipeline_config_hash and
            (base_output_dir / manifest_entry.get("output_chunks_path", "")).exists() and
            (base_output_dir / manifest_entry.get("output_parents_path", "")).exists()
            ):
            logger.info(f"Skipping source '{sid}': Manifest indicates already processed with same config. Use --force-reprocess to override.")
            doc_summary_entry.update({
                "status": "skipped_cached",
                "message": "Skipped (matched manifest config & files).",
                "chunks_generated": manifest_entry.get("chunks_generated", 0),
                "parents_generated": manifest_entry.get("parents_generated", 0)
            })
            documents.append(None) # Add a placeholder for consistency if iterating based on 'documents' length
            continue # Move to next source
        
        # If not skipped, then begin processing
        doc_summary_entry["status"] = "processing"
        doc_summary_entry["message"] = "Processing started."

        # Simplified reprocess check:
        # For a truly robust system, you'd load a manifest file that stores the config_hash
        # for each source previously processed. Here, we simulate by checking a placeholder.
        # This part requires a central manifest file. For now, it's a conceptual check.
        # If args.force_reprocess is False:
        #   Check manifest for sid AND current_source_config_hash. If match, skip.
        # else:
        #   Proceed with processing.

        # For this version, we'll let the processing run unless explicitly told NOT to reprocess
        # and we can reliably determine it's been processed. For now, always process unless --force-reprocess is false
        # AND some (non-implemented) manifest says it's ok.
        
        # --- (Existing Document Load and Normalize Logic) ---
        # This part remains mostly the same, creating `Document` objects.
        # It needs to handle raw/extracted content for structural chunkers later.
        

        full_text = ""
        pre_chunks: List[str] = [] # For CSV/Excel rows
        actual_source_path_or_url = "" # The abs path or URL that the document content was loaded from
        extracted_meta_from_loader: Dict[str, Any] = {} # For markdown front matter, etc.
        
        try:
            if source_type == "file":
                abs_path_str = Path(s_conf["abs_path"])
                # Decide if we need to load RAW content for structural markdown/html chunkers
                # This check happens *before* load_file_by_format
                strategy_to_apply = doc_specific_chunking_config.get("strategy")
                is_structural_html_md = strategy_to_apply in ["markdown_header", "html_section"]
                
                if is_structural_html_md:
                    logger.info(f"Loading RAW content for '{sid}' (format: {source_format}) due to structural chunking strategy '{strategy_to_apply}'.")
                    # Special load: raw content needed for structural splitters
                    if source_format == "md":
                        raw_content, extracted_meta_from_loader = load_md_with_frontmatter(abs_path_str)
                    elif source_format == "html":
                        raw_content = load_html_raw(abs_path_str) # Your raw html loader
                    else: # Fallback
                        raw_content = load_txt_like(abs_path_str)
                    
                    full_text = raw_content # Keep it raw for the chunker
                    pre_chunks = [] # Structural chunkers don't rely on pre_chunks
                else: # Standard loading/cleaning
                    logger.info(f"Loading file '{sid}' (format: {source_format}) for standard chunking.")
                    full_text, pre_chunks, _, extracted_meta_from_loader = load_file_by_format(
                        project_root, # Base path argument is technically not used by load_file_by_format anymore
                        s_conf["path"],  # Relative path
                        source_format,
                        s_conf.get("loader_options", {})
                    )
                actual_source_path_or_url = s_conf["abs_path"] # Store absolute path as actual source used

            elif source_type == "url":
                logger.info(f"Loading URL '{sid}': {s_conf['url']}")
                full_text, pre_chunks_from_url = load_url_content(s_conf["url"], timeout=args.timeout)
                pre_chunks.extend(pre_chunks_from_url) # Append any pre-chunks from URL (e.g., if it extracts tables)
                actual_source_path_or_url = s_conf["url"]

            else:
                logger.error(f"Unsupported type '{source_type}' for source '{sid}'. Skipping.")
                doc_summary_entry["status"] = "failed_loading"
                doc_summary_entry["message"] = f"Unsupported source type: {source_type}"
                summary_output["processed_sources_details"].append(doc_summary_entry)
                continue
            
            # For non-structural chunkers, normalize the text AFTER loading (if not already raw)
            if not is_structural_html_md:
                full_text = normalize_text(full_text)

            if not full_text.strip() and not pre_chunks:
                logger.warning(f"No text extracted or pre-chunks generated for {sid}. Skipping document.")
                doc_summary["status"] = "failed_loading"
                doc_summary["message"] = "No extractable text"
                summary_output["processed_sources_details"].append(doc_summary)
                continue
            
            # Build final metadata for the Document object
            doc_meta = {k: v for k, v in s_conf.items() if k not in {"id", "title", "path", "abs_path", "url", "chunking", "loader_options", "type", "format"}}
            doc_meta["source_id"] = sid
            doc_meta["source_type"] = source_type
            doc_meta["file_path_relative"] = s_conf.get("path") # Can be None for URL
            doc_meta["file_path_absolute"] = s_conf.get("abs_path") # Can be None for URL
            doc_meta["url"] = s_conf.get("url") # Can be None for file
            doc_meta["format"] = source_format
            doc_meta["loader_options_used"] = s_conf.get("loader_options", {})
            doc_meta.update(extracted_meta_from_loader) # Blend in any metadata from loader (e.g., MD front matter)
            doc_meta["chunking_config_for_this_source"] = doc_specific_chunking_config # Keep track of exact config
            doc_meta["source_config_hash"] = current_source_config_hash # Unique hash
            doc_meta["pipeline_config_hash"] = overall_rag_pipeline_config_hash
            
            doc = Document(
                id=sid,
                title=source_title,
                source_type=source_type,
                source=actual_source_path_or_url, # The actual path/URL content came from
                text=full_text,
                meta=doc_meta # Full processed metadata
            )

            # Attach pre-chunks to meta for chunking stage (CSV/TSV/Excel/etc)
            if pre_chunks:
                doc.meta["_pre_chunks"] = pre_chunks
                doc.meta["_structured_rows"] = True # Flag for special handling in chunking
            
            documents.append(doc)
            doc_summary_entry["status"] = "loaded"
            doc_summary_entry["message"] = "Document loaded successfully."
            summary_output["processed_sources_details"].append(doc_summary)

        except Exception as e:
            logger.error(f"Error processing source '{sid}': {e}", exc_info=True)
            doc_summary_entry["status"] = "failed_loading"
            doc_summary_entry["message"] = str(e)
            continue            
    
    documents_to_process = [d for d in documents if d is not None]
    
    summary_output["total_documents_loaded"] = len(documents_to_process)
    logger.info(f"Loaded {len(documents_to_process)} documents for processing (after filtering skips).")

    if documents_to_process:
        with final_output_docs_path.open("w", encoding="utf-8") as f:
            for d in documents_to_process:
                f.write(json.dumps(asdict(d), ensure_ascii=False) + "\n")
        logger.info(f"Wrote processed documents to {final_output_docs_path}")
    else:
        logger.warning("No documents loaded for processing. Skipping writing to output-docs.")

    if not (CHUNK_STRATEGY_REGISTRY and Chunk and BaseChunkingStrategy):
        logger.critical("Chunking components not properly loaded. Cannot perform chunking.")
        sys.exit(1)
    
    if not documents_to_process:
        logger.warning("No documents to chunk. Skipping chunking phase.")
    else:
        logger.info("Starting document chunking and writing to output files segregating children and parents...")
        
        with final_output_chunks_path.open("w", encoding="utf-8") as f_child_chunks, \
             final_output_parents_path.open("w", encoding="utf-8") as f_parent_chunks:
            
            for doc in documents_to_process:
                doc_summary_entry = next((s for s in summary_output["processed_sources_details"] if s["id"] == doc.id), None)
                if not doc_summary_entry or doc_summary_entry["status"] != "loaded":
                    logger.warning(f"Skipping chunking for document '{doc.id}' due to prior loading issue.")
                    continue

                doc_chunking_config = doc.meta.get("chunking_config_for_this_source", global_chunking_config)
                strategy_name = doc_chunking_config.get("strategy")
                strategy_params = doc_chunking_config.get(strategy_name, {})

                if not strategy_name or strategy_name not in CHUNK_STRATEGY_REGISTRY:
                    logger.error(f"No valid chunking strategy '{strategy_name}' found for document '{doc.id}'. Skipping chunking.")
                    doc_summary_entry.update({
                        "status": "failed_chunking_config",
                        "message": f"Invalid/missing chunking strategy: {strategy_name}"
                    })
                    continue
                
                processed_chunks_for_doc: List[Chunk] = []

                if doc.meta.get("_structured_rows", False) and doc.meta.get("_pre_chunks"):
                    logger.info(f"Processing structured data (pre-chunks) for '{doc.id}'.")
                    pre_chunks_data = doc.meta["_pre_chunks"]
                    for i, txt_row in enumerate(pre_chunks_data):
                        clean_txt_row = normalize_text(txt_row)
                        if not clean_txt_row: continue
                        chunk_meta = {
                            **{k: v for k, v in doc.meta.items() if not k.startswith("_")},
                            "chunk_index": i,
                            "n_chars": len(clean_txt_row),
                            "strategy": "structured_rows",
                            "structured_row_index": i,
                            "chunk_type": "child",
                            **strategy_params
                        }
                        processed_chunks_for_doc.append(Chunk(
                            id=f"{doc.id}_c{i:04d}",
                            doc_id=doc.id,
                            text=clean_txt_row,
                            metadata=safe_meta(chunk_meta)
                        ))
                else:
                    try:
                        ChunkerClass: Type[BaseChunkingStrategy] = CHUNK_STRATEGY_REGISTRY[strategy_name]
                        
                        effective_chunk_size = strategy_params.get("chunk_size", rag_config.get("chunk_size", 500))
                        effective_chunk_overlap = strategy_params.get("chunk_overlap", rag_config.get("chunk_overlap", 50))
                        
                        chunker_kwargs = {**strategy_params}
                        if strategy_name == "proposition":
                            chunker_kwargs["settings"] = settings
                            chunker_kwargs["global_llm_config"] = full_config.get("llm", {})

                        chunker_instance = ChunkerClass(
                            chunk_size=effective_chunk_size,
                            chunk_overlap=effective_chunk_overlap,
                            **chunker_kwargs
                        )
                        
                        logger.info(f"Chunking '{doc.id}' with strategy '{strategy_name}' (Size: {effective_chunk_size}, Overlap: {effective_chunk_overlap})")
                        current_doc_chunks: List[Chunk] = chunker_instance.split_text(
                            text=doc.text,
                            document_id=doc.id,
                            document_metadata={k: v for k, v in doc.meta.items() if not k.startswith("_")}
                        )
                        processed_chunks_for_doc.extend(current_doc_chunks)

                    except Exception as e:
                        logger.error(f"Failed to chunk document '{doc.id}' with strategy '{strategy_name}': {e}. Skipping.", exc_info=True)
                        doc_summary_entry.update({
                            "status": "failed_chunking",
                            "message": str(e)
                        })
                        continue
                
                current_doc_child_chunks = 0
                current_doc_parent_chunks = 0

                for chunk_obj in processed_chunks_for_doc:
                    chunk_obj.metadata["chunk_strategy_applied"] = strategy_name
                    chunk_obj.metadata["chunk_size_actual"] = len(chunk_obj.text)
                    
                    if chunk_obj.metadata.get("chunk_type") == "parent":
                        f_parent_chunks.write(json.dumps(asdict(chunk_obj), ensure_ascii=False) + "\n")
                        current_doc_parent_chunks += 1
                    else:
                        f_child_chunks.write(json.dumps(asdict(chunk_obj), ensure_ascii=False) + "\n")
                        current_doc_child_chunks += 1
                
                doc_summary_entry.update({
                     "chunks_generated": current_doc_child_chunks,
                     "parents_generated": current_doc_parent_chunks,
                     "status": "processed",
                     "message": "Chunks generated successfully."
                })
                
                summary_output["total_chunks_written"] += current_doc_child_chunks
                summary_output["total_parent_chunks_written"] += current_doc_parent_chunks

    logger.info(f"Wrote {summary_output['total_chunks_written']} child/standard chunks to {final_output_chunks_path}")
    logger.info(f"Wrote {summary_output['total_parent_chunks_written']} parent chunks to {final_output_parents_path}")
    logger.info("--- Phase A complete ---")

    # --- Final step: Save the updated manifest ---
    # Prepare manifest entries based on processed_sources_details
    for entry in summary_output["processed_sources_details"]:
        if entry["status"] == "processed" or entry["status"] == "skipped_cached":
            updated_manifest[entry["id"]] = {
                "source_processing_config_hash": entry["source_processing_config_hash"],
                "pipeline_config_hash": entry["pipeline_config_hash"],
                "output_chunks_path": str(final_output_chunks_path.relative_to(base_output_dir)),
                "output_parents_path": str(final_output_parents_path.relative_to(base_output_dir)),
                "chunks_generated": entry["chunks_generated"],
                "parents_generated": entry["parents_generated"],
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
            }
        elif entry["id"] in updated_manifest and \
            (entry["status"] == "failed_loading" or entry["status"] == "failed_chunking" or entry["status"] == "failed_chunking_config"):
            updated_manifest.pop(entry["id"], None)
            
    save_manifest(manifest_path, updated_manifest)

    # Final Summary Output to console
    print("\n--- Phase A Execution Summary ---")
    print(json.dumps(summary_output, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()